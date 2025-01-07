import cv2
import numpy as np
from PIL import Image
from pdf2image import convert_from_bytes
from docx import Document
import easyocr
import io
import streamlit as st

def is_image_blurry(file_like_object, threshold=100):
    """
    Check if an image is blurry using the variance of the Laplacian.
    :param file_like_object: File-like object containing the image.
    :param threshold: Blurriness threshold. Lower values indicate blurrier images.
    :return: True if the image is blurry, False otherwise.
    """
    file_like_object.seek(0)
    file_bytes = np.frombuffer(file_like_object.read(), np.uint8)
    image = cv2.imdecode(file_bytes, cv2.IMREAD_GRAYSCALE)
    laplacian_var = cv2.Laplacian(image, cv2.CV_64F).var()
    print(f"Laplacian Variance: {laplacian_var}")
    return laplacian_var < threshold

def extract_text_from_image(file_like_object):
    """
    Extract text from an image using EasyOCR.
    :param file_like_object: File-like object containing the image.
    :return: Extracted text.
    """
    file_like_object.seek(0)
    if is_image_blurry(file_like_object):
        raise ValueError("The image is too blurry for OCR. Please provide a clearer image.")

    file_like_object.seek(0)
    reader = easyocr.Reader(["en", "ch_sim"])  # Use ["en", "ch_sim"] for now
    final_result = reader.readtext(np.array(Image.open(file_like_object)), detail=0)
    return "\n".join(final_result)

def extract_text_from_pdf(file_like_object):
    try:
        # Attempt to extract text from PDF directly
        file_like_object.seek(0)
        from PyPDF2 import PdfReader
        reader = PdfReader(file_like_object)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        if text.strip():
            return text
    except Exception as e:
        print(f"Error reading text directly from PDF: {e}")

    return extract_text_from_pdf_ocr(file_like_object)

def extract_text_from_pdf_ocr(file_like_object):
    # Fallback to OCR on PDF images
    print("No direct text found in PDF. Using OCR on PDF images.")
    file_like_object.seek(0)
    images = convert_from_bytes(file_like_object.read())
    text = ""
    for image in images:
        temp_image = io.BytesIO()
        image.save(temp_image, format="JPEG")
        temp_image.seek(0)
        text += extract_text_from_image(temp_image) + "\n"
    return text

def extract_text_from_word(file_like_object):
    file_like_object.seek(0)
    doc = Document(file_like_object)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    # Process embedded images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            print("Embedded image detected in Word document. Applying OCR.")
            img_data = rel.target_part.blob
            image = Image.open(io.BytesIO(img_data))
            temp_image = io.BytesIO()
            image.save(temp_image, format="JPEG")
            temp_image.seek(0)
            text += extract_text_from_image(temp_image) + "\n"
    return text

def process_file(file_like_object, file_name):
    """
    Process a file and extract text based on its type.
    :param file_like_object: File-like object.
    :param file_name: Name of the uploaded file.
    :return: Extracted text.
    """
    file_extension = file_name.split(".")[-1].lower()
    if file_extension in ["jpg", "jpeg", "png", "bmp", "tiff"]:
        return extract_text_from_image(file_like_object)
    elif file_extension == "pdf":
        return extract_text_from_pdf(file_like_object)
    elif file_extension in ["docx", "doc"]:
        return extract_text_from_word(file_like_object)
    else:
        raise ValueError("Unsupported file format. Please upload an image, PDF, or Word document.")

def upload_file(uploaded_file):
    extracted_text = process_file(uploaded_file, uploaded_file.name)

    # st.write(extracted_text)

    file_path = 'reference_doc.txt'
    with open(file_path, 'w', encoding='utf-8') as file:
        # Write the string to the file
        file.write("reference document\n" + extracted_text)

    return file_path

def create_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer




# import os
# from PIL import Image
# from pdf2image import convert_from_path
# from docx import Document
# import easyocr
# import cv2
# import numpy as np
# import io

# def is_image_blurry(image_path, threshold=100):
#     """
#     Check if an image is blurry using the variance of the Laplacian.
#     :param image_path: Path to the image file.
#     :param threshold: Blurriness threshold. Lower values indicate blurrier images.
#     :return: True if the image is blurry, False otherwise.
#     """
#     image = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
#     laplacian_var = cv2.Laplacian(image, cv2.CV_64F).var()
#     print(f"Laplacian Variance: {laplacian_var}")
#     return laplacian_var < threshold

# def extract_text_from_image(image_path):
#     """
#     Extract text from an image using EasyOCR.
#     :param image_path: Path to the image file.
#     :return: Extracted text.
#     """
#     if is_image_blurry(image_path):
#         raise ValueError("The image is too blurry for OCR. Please provide a clearer image.")

   
#     reader = easyocr.Reader(["en", "ch_sim"]) #use ["en", "ch_sim"] for now
#     final_result = reader.readtext(image_path, detail=0)
#     return "\n".join(final_result)

# def extract_text_from_pdf(pdf_path):
#     try:
#         # Attempt to extract text from PDF directly
#         with open(pdf_path, "rb") as f:
#             from PyPDF2 import PdfReader
#             reader = PdfReader(f)
#             text = ""
#             for page in reader.pages:
#                 text += page.extract_text()
#         if text.strip():
#             return text
#     except Exception as e:
#         print(f"Error reading text directly from PDF: {e}")

#     return extract_text_from_pdf_ocr(pdf_path)

# def extract_text_from_pdf_ocr(pdf_path):
#     # Fallback to OCR on PDF images
#     print("No direct text found in PDF. Using OCR on PDF images.")
#     images = convert_from_path(pdf_path)
#     text = ""
#     for image in images:
#         temp_image_path = "temp_image.jpg"
#         image.save(temp_image_path)
#         text += extract_text_from_image(temp_image_path) + "\n"
#     return text

# def extract_text_from_word(word_path):
#     doc = Document(word_path)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"

#     # Process embedded images
#     for rel in doc.part.rels.values():
#         if "image" in rel.target_ref:
#             print("Embedded image detected in Word document. Applying OCR.")
#             img_data = rel.target_part.blob
#             image = Image.open(io.BytesIO(img_data))
#             temp_image_path = "temp_image.jpg"
#             image.save(temp_image_path)
#             text += extract_text_from_image(temp_image_path) + "\n"
#     return text

# def process_file(file_path):
#     """
#     Process a file and extract text based on its type.
#     :param file_path: Path to the uploaded file.
#     :return: Extracted text.
#     """
#     file_extension = file_path.split(".")[-1].lower()
#     if file_extension in ["jpg", "jpeg", "png", "bmp", "tiff"]:
#         return extract_text_from_image(file_path)
#     elif file_extension == "pdf":
#         return extract_text_from_pdf(file_path)
#     elif file_extension in ["docx", "doc"]:
#         return extract_text_from_word(file_path)
#     else:
#         raise ValueError("Unsupported file format. Please upload an image, PDF, or Word document.")


# def upload_file(uploaded_file_path):
#     extracted_text = process_file(uploaded_file_path)

#     file_path = 'reference_doc.txt'
#     with open(file_path, 'w') as file:
#         # Write the string to the file
#         file.write(extracted_text)

#     return file_path